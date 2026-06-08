#!/usr/bin/env python3
from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NOTES_HOME = Path(
    os.environ.get("VIDEO_NOTES_HOME", str(Path.home() / "Documents" / "VideoNotes"))
).expanduser()


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor(35, 35, 35)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 19, "1F4D78"),
        ("Heading 2", 14, "1F4D78"),
        ("Heading 3", 11.5, "2E74B5"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_markdown_line(doc: Document, line: str) -> None:
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
        return
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        return
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        return
    if line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
        return
    if not line.strip():
        doc.add_paragraph("")
        return
    paragraph = doc.add_paragraph()
    for idx, part in enumerate(re.split(r"(`[^`]+`)", line)):
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor(45, 45, 45)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()

    markdown = args.markdown
    if not markdown.exists():
        raise SystemExit(f"Markdown not found: {markdown}")
    out = args.out or NOTES_HOME / "exports" / f"{markdown.stem}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_doc(doc)
    for line in markdown.read_text(encoding="utf-8").splitlines():
        add_markdown_line(doc, line)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Chinese Video Notes"
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
